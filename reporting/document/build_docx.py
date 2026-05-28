"""
build_docx.py — renders the PolarisBI written report (a real consulting briefing,
prose-led) from report_data.REPORT using python-docx.

Architecture (conventions cherry-picked from McKinsey/Deloitte reports, original content):
  Cover (title block + classification + KPI strip) -> Contents -> running header/footer
  -> Executive Summary (prose) -> numbered sections (intro + bold lead-in + analysis,
  tables demoted to captioned Exhibits) -> Methodology & Data Sources.

Public entrypoint: build(report=None, out_dir=None) -> pathlib.Path to the .docx
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from reporting.report_data import REPORT

# ---------------------------------------------------------------- brand tokens
NAVY   = "0C2340"
ROYAL  = "164E96"
BLUE   = "2F6FE0"
GREEN  = "2E9E5B"
RED    = "C0392B"
AMBER  = "B7791F"
INK    = "1F2733"
GRAY   = "6B7280"
GRAYLT = "9AA3AF"
LINE   = "DCE2EC"
SOFT   = "E8F0FC"
CARD   = "F6F8FC"
WHITE  = "FFFFFF"

HEAD = "Poppins"   # display / headings (installed on host; falls back gracefully)
BODY = "Inter"     # body text

CONTENT_W = 9360   # US Letter, 1" margins, in DXA/twips (8.5" - 2")


# ----------------------------------------------------------------- low-level helpers
def _set_font(run, font=BODY, size=10.5, bold=False, italic=False, color=INK,
              small_caps=False, spacing=None):
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), font)
    if small_caps:
        sc = OxmlElement("w:smallCaps")
        rpr.append(sc)
    if spacing is not None:  # letter-spacing in twentieths of a point
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:val"), str(spacing))
        rpr.append(sp)


def _shade(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcpr.append(shd)


def _cell_margins(cell, top=60, bottom=60, left=110, right=110):
    tcpr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right),
                     ("left", left), ("right", right)):
        e = OxmlElement(f"w:{tag}")
        e.set(qn("w:w"), str(val))
        e.set(qn("w:type"), "dxa")
        m.append(e)
    tcpr.append(m)


def _no_table_borders(table):
    tblpr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "none")
        borders.append(e)
    tblpr.append(borders)


def _row_bottom_border(cell, color=LINE, size=4):
    tcpr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), str(size))
    b.set(qn("w:color"), color)
    borders.append(b)
    tcpr.append(borders)


def _para_rule(paragraph, color=NAVY, size=8, space=2, edge="bottom"):
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    b = OxmlElement(f"w:{edge}")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), str(size))
    b.set(qn("w:space"), str(space))
    b.set(qn("w:color"), color)
    pbdr.append(b)


def _set_col_widths(table, widths):
    table.autofit = False
    table.allow_autofit = False
    for row in table.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = Twips(w)
    # also set tblGrid so Word/LibreOffice honor widths
    tbl = table._tbl
    grid = tbl.find(qn("w:tblGrid"))
    if grid is not None:
        tbl.remove(grid)
    grid = OxmlElement("w:tblGrid")
    for w in widths:
        c = OxmlElement("w:gridCol")
        c.set(qn("w:w"), str(w))
        grid.append(c)
    tbl.tblPr.addnext(grid)


def _keep_with_next(paragraph):
    ppr = paragraph._p.get_or_add_pPr()
    k = OxmlElement("w:keepNext")
    ppr.append(k)


def _no_split_rows(table):
    for row in table.rows:
        trpr = row._tr.get_or_add_trPr()
        cant = OxmlElement("w:cantSplit")
        trpr.append(cant)


def _add_page_number(paragraph):
    run = paragraph.add_run()
    fb = OxmlElement("w:fldChar"); fb.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = " PAGE "
    fe = OxmlElement("w:fldChar"); fe.set(qn("w:fldCharType"), "end")
    run._r.append(fb); run._r.append(it); run._r.append(fe)
    _set_font(run, BODY, 8.5, color=GRAY)


def _blank(paragraph_factory, pts=6):
    p = paragraph_factory()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run("")
    _set_font(r, BODY, pts)
    return p


# ----------------------------------------------------------------- page geometry
def _setup_page(section):
    section.page_width = Twips(12240)
    section.page_height = Twips(15840)
    section.top_margin = Twips(1180)
    section.bottom_margin = Twips(1120)
    section.left_margin = Twips(1440)
    section.right_margin = Twips(1440)
    section.header_distance = Twips(720)
    section.footer_distance = Twips(620)


def _build_header_footer(section, meta):
    section.different_first_page_header_footer = True  # cover stays clean

    # ---- running header (pages 2+)
    hdr = section.header
    hp = hdr.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = hp.add_run(meta["running_title"])
    _set_font(r, BODY, 8, bold=False, color=GRAYLT, small_caps=True, spacing=20)
    _para_rule(hp, color=LINE, size=4, space=4, edge="bottom")

    # ---- running footer (pages 2+): confidentiality (left) + page number (right)
    ftr = section.footer
    fp = ftr.paragraphs[0]
    fp.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    _para_rule(fp, color=LINE, size=4, space=4, edge="top")
    r = fp.add_run(meta["footer"])
    _set_font(r, BODY, 8, color=GRAY)
    rt = fp.add_run("\t")
    _set_font(rt, BODY, 8, color=GRAY)
    pre = fp.add_run("Page ")
    _set_font(pre, BODY, 8.5, color=GRAY)
    _add_page_number(fp)

    # keep first-page header/footer empty
    section.first_page_header.paragraphs[0].text = ""
    section.first_page_footer.paragraphs[0].text = ""


# ----------------------------------------------------------------- cover
def _build_cover(doc, meta, kpis):
    for _ in range(2):
        _blank(doc.add_paragraph, 6)

    p = doc.add_paragraph()
    r = p.add_run(meta["eyebrow"].upper())
    _set_font(r, BODY, 9.5, bold=True, color=ROYAL, small_caps=True, spacing=40)
    p.paragraph_format.space_after = Pt(14)

    p = doc.add_paragraph()
    r = p.add_run(meta["title"])
    _set_font(r, HEAD, 30, bold=True, color=NAVY)
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    _para_rule(p, color=NAVY, size=18, space=6, edge="bottom")
    p.paragraph_format.space_after = Pt(12)

    p = doc.add_paragraph()
    r = p.add_run(meta["subtitle"])
    _set_font(r, BODY, 13, color=GRAY)
    p.paragraph_format.space_after = Pt(30)

    # KPI strip — 4 columns, borderless (skip if none provided)
    if kpis:
        tbl = doc.add_table(rows=1, cols=4)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        _no_table_borders(tbl)
        _set_col_widths(tbl, [CONTENT_W // 4] * 4)
        for i, k in enumerate(kpis):
            cell = tbl.rows[0].cells[i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            _cell_margins(cell, top=40, bottom=40, left=0, right=140)
            pl = cell.paragraphs[0]
            rr = pl.add_run(k["label"].upper())
            _set_font(rr, BODY, 7.5, bold=True, color=GRAYLT, small_caps=True, spacing=20)
            pl.paragraph_format.space_after = Pt(1)
            pv = cell.add_paragraph()
            rr = pv.add_run(k["value"])
            _set_font(rr, HEAD, 17, bold=True, color=NAVY)
            pv.paragraph_format.space_after = Pt(1)
            pd = cell.add_paragraph()
            rr = pd.add_run(k["delta"])
            _set_font(rr, BODY, 9, bold=True, color=GREEN)

    # push author block toward the bottom
    for _ in range(8):
        _blank(doc.add_paragraph, 6)

    p = doc.add_paragraph()
    _para_rule(p, color=LINE, size=6, space=6, edge="bottom")
    p.paragraph_format.space_after = Pt(8)

    p = doc.add_paragraph()
    r = p.add_run(meta["classification"])
    _set_font(r, BODY, 9.5, italic=True, color=GRAY)
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    r = p.add_run(f'{meta["author"]}   ·   {meta["date"]}')
    _set_font(r, BODY, 10, color=INK)

    doc.add_page_break()


# ----------------------------------------------------------------- contents
def _build_contents(doc, sections, include_methodology=True):
    p = doc.add_paragraph()
    r = p.add_run("Contents")
    _set_font(r, HEAD, 18, bold=True, color=NAVY)
    _para_rule(p, color=NAVY, size=8, space=4, edge="bottom")
    p.paragraph_format.space_after = Pt(16)

    entries = [("", "Executive Summary")]
    for s in sections:
        entries.append((s["num"], s["title"]))
    if include_methodology:
        entries.append(("", "Methodology & Data Sources"))

    for num, title in entries:
        line = doc.add_paragraph()
        line.paragraph_format.space_after = Pt(9)
        if num:
            rn = line.add_run(f"{num}    ")
            _set_font(rn, HEAD, 11.5, bold=True, color=ROYAL)
        rt = line.add_run(title)
        _set_font(rt, BODY, 11.5, color=INK)

    doc.add_page_break()


# ----------------------------------------------------------------- exec summary
def _build_exec_summary(doc, paras):
    h = doc.add_paragraph()
    r = h.add_run("Executive Summary")
    _set_font(r, HEAD, 16, bold=True, color=NAVY)
    _para_rule(h, color=NAVY, size=8, space=4, edge="bottom")
    h.paragraph_format.space_before = Pt(2)
    h.paragraph_format.space_after = Pt(12)

    for para in paras:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(11)
        p.paragraph_format.line_spacing = 1.28
        r = p.add_run(para)
        _set_font(r, BODY, 10.5, color=INK)


# ----------------------------------------------------------------- section pieces
def _section_heading(doc, num, title):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(18)
    h.paragraph_format.space_after = Pt(10)
    _keep_with_next(h)
    rn = h.add_run(f"{num}.  ")
    _set_font(rn, HEAD, 15, bold=True, color=ROYAL)
    rt = h.add_run(title)
    _set_font(rt, HEAD, 15, bold=True, color=NAVY)
    _para_rule(h, color=NAVY, size=6, space=4, edge="bottom")


def _intro_para(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(11)
    p.paragraph_format.line_spacing = 1.28
    r = p.add_run(text)
    _set_font(r, BODY, 10.5, color=INK)


def _lead_in_para(doc, lead, body):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.28
    # colour severity tags if present
    color = ROYAL
    if lead.startswith("[HIGH]"):
        color = RED
    elif lead.startswith("[MEDIUM]"):
        color = AMBER
    elif lead.startswith("[LOW]"):
        color = GREEN
    rl = p.add_run(lead + " ")
    _set_font(rl, BODY, 10.5, bold=True, color=color)
    rb = p.add_run(body)
    _set_font(rb, BODY, 10.5, color=INK)


def _caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(5)
    _keep_with_next(p)
    r = p.add_run(text)
    _set_font(r, BODY, 8.5, bold=True, color=GRAY, small_caps=True, spacing=15)


def _source_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("Source: " + text)
    _set_font(r, BODY, 8, italic=True, color=GRAYLT)


# ----------------------------------------------------------------- exhibits/tables
def _exhibit_trend(doc, ex):
    _caption(doc, ex["caption"])
    n = len(ex["labels"])
    tbl = doc.add_table(rows=2, cols=n)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_col_widths(tbl, [CONTENT_W // n] * n)
    for i, lab in enumerate(ex["labels"]):
        c = tbl.rows[0].cells[i]
        _shade(c, NAVY); _cell_margins(c)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        pp = c.paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = pp.add_run(lab); _set_font(r, BODY, 9.5, bold=True, color=WHITE)
    for i, val in enumerate(ex["values"]):
        c = tbl.rows[1].cells[i]
        _shade(c, SOFT); _cell_margins(c, top=80, bottom=80)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        pp = c.paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = pp.add_run(val); _set_font(r, HEAD, 14, bold=True, color=NAVY)
    _no_split_rows(tbl)
    _source_note(doc, ex["source"])


def _exhibit_simple(doc, t):
    """Generic table: first row = header (navy), optional highlight_row (data index)."""
    _caption(doc, t["caption"])
    header = t["header"]
    rows = t["rows"]
    ncol = len(header)
    tbl = doc.add_table(rows=1 + len(rows), cols=ncol)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # column widths: first col wide if it's a label table; numbers get even share
    if t.get("caption", "").startswith("Exhibit 4") or header[0] == "#":
        widths = [560, 3200, 1320, 1080, 1080, 1080, 1040]
        widths = widths[:ncol]
        diff = CONTENT_W - sum(widths); widths[1] += diff
    elif header[0] in ("Channel", "Metric"):
        first = 2600
        rest = (CONTENT_W - first) // (ncol - 1)
        widths = [first] + [rest] * (ncol - 1)
        widths[-1] += CONTENT_W - sum(widths)
    else:
        even = CONTENT_W // ncol
        widths = [even] * ncol
        widths[-1] += CONTENT_W - sum(widths)
    _set_col_widths(tbl, widths)

    numeric_from = 1  # right-align columns after the first
    for i, htxt in enumerate(header):
        c = tbl.rows[0].cells[i]
        _shade(c, NAVY); _cell_margins(c)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        pp = c.paragraphs[0]
        pp.alignment = WD_ALIGN_PARAGRAPH.LEFT if i < numeric_from else WD_ALIGN_PARAGRAPH.RIGHT
        r = pp.add_run(str(htxt)); _set_font(r, BODY, 9, bold=True, color=WHITE)

    hi = t.get("highlight_row", -1)
    for ridx, row in enumerate(rows):
        zebra = CARD if (ridx % 2 == 1) else WHITE
        fill = SOFT if ridx == hi else zebra
        for i, val in enumerate(row):
            c = tbl.rows[ridx + 1].cells[i]
            _shade(c, fill); _cell_margins(c)
            _row_bottom_border(c, color=LINE, size=4)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            pp = c.paragraphs[0]
            pp.alignment = WD_ALIGN_PARAGRAPH.LEFT if i < numeric_from else WD_ALIGN_PARAGRAPH.RIGHT
            is_hi = ridx == hi
            r = pp.add_run(str(val))
            _set_font(r, BODY, 9.5, bold=is_hi, color=(ROYAL if is_hi else INK))
    _no_split_rows(tbl)
    if t.get("source"):
        _source_note(doc, t["source"])
    else:
        sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(12)


def _exhibit_comparison(doc, ex):
    _caption(doc, ex["caption"])
    cols = ex["cols"]          # company names
    rows = ex["rows"]          # [label, (val,rank)*4]
    ncol = 1 + len(cols)
    tbl = doc.add_table(rows=1 + len(rows), cols=ncol)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    first = 2300
    rest = (CONTENT_W - first) // len(cols)
    widths = [first] + [rest] * len(cols)
    widths[-1] += CONTENT_W - sum(widths)
    _set_col_widths(tbl, widths)

    # header row
    hc = tbl.rows[0].cells[0]
    _shade(hc, WHITE); _cell_margins(hc)
    pp = hc.paragraphs[0]; r = pp.add_run("Metric")
    _set_font(r, BODY, 9, bold=True, color=GRAY, small_caps=True, spacing=15)
    for j, name in enumerate(cols):
        c = tbl.rows[0].cells[j + 1]
        _shade(c, NAVY if j == 0 else CARD); _cell_margins(c)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        pp = c.paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = pp.add_run(name)
        _set_font(r, BODY, 9.5, bold=True, color=(WHITE if j == 0 else NAVY))

    for ridx, row in enumerate(rows):
        label = row[0]
        c = tbl.rows[ridx + 1].cells[0]
        _shade(c, WHITE); _cell_margins(c); _row_bottom_border(c)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        pp = c.paragraphs[0]; r = pp.add_run(label)
        _set_font(r, BODY, 9.5, color=GRAY)
        for j in range(len(cols)):
            val, rank = row[j + 1]
            c = tbl.rows[ridx + 1].cells[j + 1]
            _shade(c, SOFT if j == 0 else WHITE); _cell_margins(c); _row_bottom_border(c)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            pp = c.paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = pp.add_run(val + "  ")
            _set_font(r, BODY, 9.5, bold=(j == 0), color=(ROYAL if j == 0 else INK))
            rr = pp.add_run(rank)
            _set_font(rr, BODY, 8, color=GRAYLT)
    _no_split_rows(tbl)
    _source_note(doc, ex["source"])


# ----------------------------------------------------------------- recs & methodology
def _build_recommendations(doc, section):
    _section_heading(doc, section["num"], section["title"])
    _intro_para(doc, section["intro"])
    for i, rec in enumerate(section["recs"], start=1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        _keep_with_next(p)
        rn = p.add_run(f"Priority {i} — ")
        _set_font(rn, BODY, 10.5, bold=True, color=ROYAL)
        rt = p.add_run(rec["title"])
        _set_font(rt, BODY, 10.5, bold=True, color=NAVY)

        b = doc.add_paragraph()
        b.paragraph_format.space_after = Pt(4)
        b.paragraph_format.line_spacing = 1.26
        r = b.add_run(rec["body"])
        _set_font(r, BODY, 10.5, color=INK)

        m = doc.add_paragraph()
        m.paragraph_format.space_after = Pt(10)
        for lbl, val in (("Owner", rec["owner"]), ("Timeline", rec["timeline"]),
                         ("Target", rec["metric"])):
            rk = m.add_run(f"{lbl}: ")
            _set_font(rk, BODY, 9, bold=True, color=GRAY, small_caps=True, spacing=10)
            rv = m.add_run(val + "     ")
            _set_font(rv, BODY, 9.5, color=INK)


def _build_methodology(doc, meth):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(18)
    h.paragraph_format.space_after = Pt(10)
    _keep_with_next(h)
    r = h.add_run("Methodology & Data Sources")
    _set_font(r, HEAD, 15, bold=True, color=NAVY)
    _para_rule(h, color=NAVY, size=6, space=4, edge="bottom")

    _intro_para(doc, meth["intro"])
    for i, src in enumerate(meth["sources"], start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(14)
        p.paragraph_format.space_after = Pt(5)
        rn = p.add_run(f"{i}.  ")
        _set_font(rn, BODY, 10, bold=True, color=ROYAL)
        rt = p.add_run(src)
        _set_font(rt, BODY, 10, color=INK)


# ----------------------------------------------------------------- main build
def build(report=None, out_dir=None):
    """Render the report. If `report` is None, use the static REPORT (BRI Life
    briefing). Pass a dict in the same schema to render a dynamic/session report."""
    R = report if report is not None else REPORT
    out_dir = Path(out_dir) if out_dir else (Path(__file__).resolve().parent / "out")
    out_dir.mkdir(parents=True, exist_ok=True)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = BODY
    style.font.size = Pt(10.5)
    style.font.color.rgb = RGBColor.from_string(INK)

    section = doc.sections[0]
    _setup_page(section)

    meta = R["meta"]
    _build_cover(doc, meta, R.get("kpis", []))
    _build_header_footer(section, meta)   # applies to pages 2+
    _build_contents(doc, R["sections"], include_methodology=("methodology" in R))
    _build_exec_summary(doc, R["exec_summary"])

    for s in R["sections"]:
        if "recs" in s:
            _build_recommendations(doc, s)
            continue
        _section_heading(doc, s["num"], s["title"])
        if s.get("intro"):
            _intro_para(doc, s["intro"])
        for lead, body in s.get("points", []):
            _lead_in_para(doc, lead, body)
        if "exhibit" in s:
            ex = s["exhibit"]
            if ex["type"] == "trend":
                _exhibit_trend(doc, ex)
            elif ex["type"] == "comparison":
                _exhibit_comparison(doc, ex)
            else:  # table / channel / generic
                _exhibit_simple(doc, ex)
        if "table" in s:
            _exhibit_simple(doc, s["table"])

    if "methodology" in R:
        _build_methodology(doc, R["methodology"])

    fname = meta.get("filename", "PolarisBI-BRILife-FY2024-Report.docx")
    out_path = out_dir / fname
    doc.save(str(out_path))
    print(f"WROTE {out_path}")
    return out_path


if __name__ == "__main__":
    p = build()
    print("WROTE", p)
